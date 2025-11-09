import os, io, re, json, time, uuid, base64, zipfile, random, string, textwrap
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse
from datetime import datetime, timezone

import requests
import boto3
import nltk
import streamlit as st
from collections import OrderedDict
from dotenv import load_dotenv
from openai import AzureOpenAI

# Azure Speech SDK (for neural voices)
import azure.cognitiveservices.speech as speechsdk

# Azure Document Intelligence
try:
    from azure.ai.formrecognizer import DocumentAnalysisClient
    from azure.core.credentials import AzureKeyCredential
    DI_AVAILABLE = True
except ImportError:
    DI_AVAILABLE = False

# PDF/DOCX extraction (fallback)
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# =========================
# Base Config & Utilities
# =========================
load_dotenv()

# NLTK once
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# ---------- Constants ----------
DEFAULT_COVER_URL = "https://media.suvichaar.org/upload/covers/default_news.png"
DEFAULT_SLIDE_IMAGE_URL = "https://media.suvichaar.org/upload/covers/default_news_slide.png"
DEFAULT_CTA_AUDIO = "https://cdn.suvichaar.org/media/tts_cta_default.mp3"
SLIDE_CHAR_LIMITS = {
    1: 80,
    2: 500,
    3: 450,
    4: 250,
    5: 200,
    "default": 200,
}

# ---- Azure OpenAI Client (for text/gen) ----
client = AzureOpenAI(
    azure_endpoint= st.secrets["azure_api"]["AZURE_OPENAI_ENDPOINT"],
    api_key= st.secrets["azure_api"]["AZURE_OPENAI_API_KEY"],
    api_version="2025-01-01-preview"
)

# ---- Azure Speech (Neural TTS) ----
AZURE_SPEECH_KEY   = st.secrets["azure"]["AZURE_API_KEY"]
AZURE_SPEECH_REGION = st.secrets["azure"].get("AZURE_REGION", "eastus")  # ensure region matches your resource

# ---- Azure Document Intelligence ----
AZURE_DI_ENDPOINT = st.secrets["azure_di"]["AZURE_DI_ENDPOINT"]
AZURE_DI_KEY = st.secrets["azure_di"]["AZURE_DI_KEY"]

# Initialize Document Intelligence client
if DI_AVAILABLE:
    try:
        di_client = DocumentAnalysisClient(
            endpoint=AZURE_DI_ENDPOINT,
            credential=AzureKeyCredential(AZURE_DI_KEY)
        )
    except Exception as e:
        di_client = None
        print(f"⚠️ Document Intelligence client initialization failed: {e}")
else:
    di_client = None

# ---- AWS ----
AWS_ACCESS_KEY = st.secrets["aws"]["AWS_ACCESS_KEY"]
AWS_SECRET_KEY = st.secrets["aws"]["AWS_SECRET_KEY"]
AWS_REGION     = st.secrets["aws"]["AWS_REGION"]
AWS_BUCKET     = st.secrets["aws"]["AWS_BUCKET"]        # unified bucket usage
S3_PREFIX      = st.secrets["aws"].get("S3_PREFIX", "media/")
CDN_BASE       = st.secrets["aws"]["CDN_BASE"]
CDN_PREFIX_MEDIA = "https://media.suvichaar.org/"

s3_client = boto3.client(
    "s3",
    aws_access_key_id     = AWS_ACCESS_KEY,
    aws_secret_access_key = AWS_SECRET_KEY,
    region_name           = AWS_REGION,
)

# ---- Voice Options (Azure Neural) ----
def pick_voice_for_language(lang_code: str, default_voice: str) -> str:
    """Map detected language → Azure voice name."""
    if not lang_code:
        return default_voice
    l = lang_code.lower()
    if l.startswith("hi"):
        return "hi-IN-AaravNeural"
    if l.startswith("en-in"):
        return "en-IN-NeerjaNeural"
    if l.startswith("en"):
        return "en-IN-AaravNeural"
    if l.startswith("bn"):
        return "bn-IN-BashkarNeural"
    if l.startswith("ta"):
        return "ta-IN-PallaviNeural"
    if l.startswith("te"):
        return "te-IN-ShrutiNeural"
    if l.startswith("mr"):
        return "mr-IN-AarohiNeural"
    if l.startswith("gu"):
        return "gu-IN-DhwaniNeural"
    if l.startswith("kn"):
        return "kn-IN-SapnaNeural"
    if l.startswith("pa"):
        return "pa-IN-GeetikaNeural"
    return default_voice

voice_options = {
    "1": "en-IN-AaravNeural",
    "2": "hi-IN-AaravNeural",
    "3": "bn-IN-BashkarNeural",
    "4": "ta-IN-PallaviNeural",
    "5": "te-IN-ShrutiNeural",
    "6": "mr-IN-AarohiNeural",
    "7": "gu-IN-DhwaniNeural",
    "8": "kn-IN-SapnaNeural",
    "9": "pa-IN-GeetikaNeural"
}
# Slug and URL generator (matching JavaScript Canurl function)
def generate_slug_and_urls(title):
    if not title or not isinstance(title, str):
        raise ValueError("Invalid title: Title must be a non-empty string.")
    
    # Step 1: Slugify the title (matching JavaScript regex)
    slug = title.lower()
    slug = re.sub(r'\s+', '-', slug)  # Replace spaces with hyphens
    slug = re.sub(r'[^a-z0-9-]', '', slug)  # Remove non-alphanumeric characters (except hyphens)
    slug = re.sub(r'^-+|-+$', '', slug)  # Remove leading or trailing hyphens
    
    # Step 2: Generate a Nano ID (matching JavaScript Canurl function)
    alphabet = 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789_-'
    size = 10
    nano_id = ''.join(random.choices(alphabet, k=size))
    nano = nano_id + "_G"
    
    slug_nano = f"{slug}_{nano}"
    
    # Step 3: Return URLs (matching JavaScript Canurl return format)
    # canurl = https://suvichaar.org/stories/slug_nano
    # canurl1 = https://stories.suvichaar.org/slug_nano.html
    return nano, slug_nano, f"https://suvichaar.org/stories/{slug_nano}", f"https://stories.suvichaar.org/{slug_nano}.html"

# === Utility Functions ===
def extract_article(url):
    import newspaper
    from newspaper import Article

    try:
        article = Article(url)
        article.download()
        article.parse()

        try:
            article.nlp()
        except:
            pass  # Some articles may not support NLP extraction

        # Fallbacks for missing fields
        title = article.title or "Untitled Article"
        text = article.text or "No article content available."
        summary = article.summary or text[:300]

        return title.strip(), summary.strip(), text.strip()

    except Exception as e:
        st.error(f"❌ Failed to extract article from URL. Error: {str(e)}")
        return "Untitled Article", "No summary available.", "No article content available."


def extract_from_pdf(file):
    """Extract text from PDF file using Azure Document Intelligence"""
    # Try Azure Document Intelligence first
    if DI_AVAILABLE and di_client:
        try:
            # Read file content
            file.seek(0)
            file_bytes = file.read()
            
            # Analyze document using prebuilt-read model
            poller = di_client.begin_analyze_document(
                model_id="prebuilt-read",
                document=file_bytes
            )
            result = poller.result()
            
            # Extract text - try content property first (most reliable)
            text = ""
            if hasattr(result, 'content') and result.content:
                text = result.content
            else:
                # Fallback: extract from pages
                for page in result.pages:
                    for line in page.lines:
                        text += line.content + "\n"
                
                # If still no text, try paragraphs
                if not text.strip() and hasattr(result, 'pages'):
                    for page in result.pages:
                        if hasattr(page, 'paragraphs'):
                            for paragraph in page.paragraphs:
                                text += paragraph.content + "\n"
            
            if not text.strip():
                return None, None, None
            
            # Use first line as title, first 300 chars as summary
            lines = text.strip().split('\n')
            title = lines[0] if lines else "Untitled Document"
            summary = text[:300].strip()
            
            return title.strip(), summary.strip(), text.strip()
        except Exception as e:
            st.warning(f"⚠️ Azure Document Intelligence failed: {str(e)}. Trying fallback...")
            # Fall through to fallback method
    
    # Fallback to PyPDF2
    if PDF_AVAILABLE:
        try:
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                return None, None, None
            
            lines = text.strip().split('\n')
            title = lines[0] if lines else "Untitled Document"
            summary = text[:300].strip()
            
            return title.strip(), summary.strip(), text.strip()
        except Exception as e:
            st.error(f"❌ Failed to extract text from PDF. Error: {str(e)}")
            return None, None, None
    else:
        st.error("❌ No PDF extraction method available. Please install azure-ai-formrecognizer or PyPDF2")
        return None, None, None


def extract_from_docx(file):
    """Extract text from DOCX file using Azure Document Intelligence"""
    # Try Azure Document Intelligence first
    if DI_AVAILABLE and di_client:
        try:
            # Read file content
            file.seek(0)
            file_bytes = file.read()
            
            # Analyze document using prebuilt-read model
            poller = di_client.begin_analyze_document(
                model_id="prebuilt-read",
                document=file_bytes
            )
            result = poller.result()
            
            # Extract text - try content property first (most reliable)
            text = ""
            if hasattr(result, 'content') and result.content:
                text = result.content
            else:
                # Fallback: extract from pages
                for page in result.pages:
                    for line in page.lines:
                        text += line.content + "\n"
                
                # If still no text, try paragraphs
                if not text.strip() and hasattr(result, 'pages'):
                    for page in result.pages:
                        if hasattr(page, 'paragraphs'):
                            for paragraph in page.paragraphs:
                                text += paragraph.content + "\n"
            
            if not text.strip():
                return None, None, None
            
            # Use first line as title, first 300 chars as summary
            lines = text.strip().split('\n')
            title = lines[0] if lines else "Untitled Document"
            summary = text[:300].strip()
            
            return title.strip(), summary.strip(), text.strip()
        except Exception as e:
            st.warning(f"⚠️ Azure Document Intelligence failed: {str(e)}. Trying fallback...")
            # Fall through to fallback method
    
    # Fallback to python-docx
    if DOCX_AVAILABLE:
        try:
            file.seek(0)
            doc = Document(file)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            
            if not text.strip():
                return None, None, None
            
            title = paragraphs[0] if paragraphs else "Untitled Document"
            summary = text[:300].strip()
            
            return title.strip(), summary.strip(), text.strip()
        except Exception as e:
            st.error(f"❌ Failed to extract text from DOCX. Error: {str(e)}")
            return None, None, None
    else:
        st.error("❌ No DOCX extraction method available. Please install azure-ai-formrecognizer or python-docx")
        return None, None, None


def extract_text_from_image(image_file):
    """Extract text from image using Azure Document Intelligence OCR"""
    if not DI_AVAILABLE or not di_client:
        return None
    
    try:
        # Read image file
        image_file.seek(0)
        image_bytes = image_file.read()
        
        # Analyze document (OCR) using prebuilt-read model
        poller = di_client.begin_analyze_document(
            model_id="prebuilt-read",
            document=image_bytes
        )
        result = poller.result()
        
        # Extract text - try content property first (most reliable)
        text = ""
        if hasattr(result, 'content') and result.content:
            text = result.content
        else:
            # Fallback: extract from pages
            for page in result.pages:
                for line in page.lines:
                    text += line.content + "\n"
            
            # If still no text, try paragraphs
            if not text.strip() and hasattr(result, 'pages'):
                for page in result.pages:
                    if hasattr(page, 'paragraphs'):
                        for paragraph in page.paragraphs:
                            text += paragraph.content + "\n"
        
        return text.strip() if text.strip() else None
    except Exception as e:
        st.warning(f"⚠️ OCR extraction failed: {str(e)}")
        return None


def upload_image_to_s3(image_file, bucket_name="suvichaarapp"):
    """Upload image to S3 bucket and return CDN URL (original, no resize)"""
    try:
        # Generate unique filename
        file_ext = os.path.splitext(image_file.name)[1] or ".jpg"
        filename = f"upload_{uuid.uuid4().hex}{file_ext}"
        s3_key = f"{S3_PREFIX}{filename}"
        
        # Read image file
        image_file.seek(0)
        image_bytes = image_file.read()
        
        # Upload to S3
        s3_client.put_object(
            Bucket=bucket_name,
            Key=s3_key,
            Body=image_bytes,
            ContentType=image_file.type or "image/jpeg"
        )
        
        # Return CDN URL (original, no resize)
        cdn_url = f"{CDN_BASE}{s3_key}"
        return cdn_url, s3_key
    except Exception as e:
        st.error(f"❌ Failed to upload image to S3. Error: {str(e)}")
        return None, None


def get_resized_image_url(s3_key, width=412, height=618, bucket_name="suvichaarapp"):
    """Generate resized image URL using CloudFront filters for slide backgrounds"""
    try:
        # Create resize template for CloudFront image transformation
        template = {
            "bucket": bucket_name,
            "key": s3_key,
            "edits": {
                "resize": {
                    "width": width,
                    "height": height,
                    "fit": "cover"  # Cover fit to maintain aspect ratio and fill the dimensions
                }
            }
        }
        encoded = base64.urlsafe_b64encode(json.dumps(template).encode()).decode()
        resized_url = f"{CDN_PREFIX_MEDIA}{encoded}"
        return resized_url
    except Exception as e:
        st.warning(f"⚠️ Failed to generate resized image URL: {str(e)}")
        return None


def get_sentiment(text):
    from textblob import TextBlob

    if not text or not text.strip():
        return "neutral"  # default for empty input

    # Clean and analyze
    clean_text = text.strip().replace("\n", " ")
    polarity = TextBlob(clean_text).sentiment.polarity

    if polarity > 0.2:
        return "positive"
    elif polarity < -0.2:
        return "negative"
    else:
        return "neutral"

def detect_category_and_subcategory(text, content_language="English"):
    import json

    if not text or len(text.strip()) < 50:
        return {
            "category": "Unknown",
            "subcategory": "General",
            "emotion": "Neutral"
        }

    # Prompt construction based on language
    if content_language == "Hindi":
        prompt = f"""
आप एक समाचार विश्लेषण विशेषज्ञ हैं।

इस समाचार लेख का विश्लेषण करें और नीचे तीन बातें बताएं:

1. category (श्रेणी)
2. subcategory (उपश्रेणी)
3. emotion (भावना)

लेख:
\"\"\"{text[:3000]}\"\"\"

जवाब केवल JSON में दें:
{{
  "category": "...",
  "subcategory": "...",
  "emotion": "..."
}}
"""
    else:
        prompt = f"""
You are an expert news analyst.

Analyze the following news article and return:

1. category
2. subcategory
3. emotion

Article:
\"\"\"{text[:3000]}\"\"\"

Return ONLY as JSON:
{{
  "category": "...",
  "subcategory": "...",
  "emotion": "..."
}}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-5-chat",
            messages=[
                {"role": "system", "content": "Classify the news into category, subcategory, and emotion."},
                {"role": "user", "content": prompt.strip()}
            ],
            max_tokens=150
        )

        content = response.choices[0].message.content.strip()
        content = content.strip("```json").strip("```").strip()

        result = json.loads(content)

        if all(k in result for k in ["category", "subcategory", "emotion"]):
            return result

    except Exception as e:
        print("❌ Category detection failed:", e)

    return {
        "category": "Unknown",
        "subcategory": "General",
        "emotion": "Neutral"
    }


def title_script_generator(
    category,
    subcategory,
    emotion,
    article_text,
    content_language="English",
    character_sketch=None,
    middle_count=5,
    slide_char_limits=None,
):
    if not character_sketch:
        character_sketch = (
            f"Polaris is a sincere and articulate {content_language} news anchor. "
            "They present facts clearly, concisely, and warmly, connecting deeply with their audience."
        )

    if slide_char_limits is None:
        slide_char_limits = SLIDE_CHAR_LIMITS

    default_limit = slide_char_limits.get("default", 200)
    guidance_map = {
        2: "detail the core development with precise names, locations, and the headline claim.",
        3: "explain earlier context, build-up, or precedent events that shaped the story.",
        4: "highlight supporting evidence—quotes, data points, documents, or eyewitness accounts.",
        5: "capture reactions from officials, experts, or the public and note immediate fallout.",
        6: "examine broader implications such as geopolitical, economic, or social impact.",
        7: "surface remaining questions, unresolved angles, or investigative threads still open.",
    }

    guidance_lines = []
    for story_slide in range(2, middle_count + 2):
        limit = slide_char_limits.get(story_slide, default_limit)
        description = guidance_map.get(
            story_slide,
            "add further factual detail, supporting evidence, or expert insight while staying concise."
        )
        guidance_lines.append(
            f"- Content Slide {story_slide - 1} (≤ {limit} characters): {description}"
        )

    guidance_text = "\n".join(guidance_lines) or "- Provide factual narrative for each slide."
    language_clause = (
        "Write all slide titles and prompts in Hindi (Devanagari script)."
        if content_language == "Hindi"
        else "Write all slide titles and prompts in English, even if the article text is in another language."
    )

    system_prompt = f"""
Create an engaging Google Web Story based on the news article provided below.

Objectives:
- Extract the key highlights, timelines, verified facts, and impactful quotes.
- Summarize the complete story visually across {middle_count} slides (target range 8–10 when article length allows).
- Keep the tone informative, balanced, and visually compelling.
- Provide slide-wise captions and background image suggestions that align with each phase of the story.
- Maintain chronological flow: introduction → build-up → evidence → reactions → implications → outlook.
- Avoid repetition; each slide must surface fresh details pulled from different portions of the article.

Language requirements:
- {language_clause}
- All fields must be written in {content_language}.

Return JSON strictly in this format:
{{
  "slides": [
    {{
      "title": "<concise slide caption (≤ 90 characters)>",
      "summary": "<two or three sentences covering the facts for narration>",
      "image_prompt": "<background or visual suggestion relevant to this slide>"
    }},
    ...
  ]
}}
"""

    user_prompt = f"""
Category: {category}
Subcategory: {subcategory}
Emotion: {emotion}

Article:
\"\"\"{article_text[:3000]}\"\"\"
"""

    response = client.chat.completions.create(
        model="gpt-5-chat",
        messages=[
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": user_prompt.strip()}
        ]
    )

    content = response.choices[0].message.content.strip()
    content = content.strip("```json").strip("```").strip()

    try:
        slides_raw = json.loads(content)["slides"]
    except:
        return {"category": category, "subcategory": subcategory, "emotion": emotion, "slides": []}

    # 🔹 Generate Slide 1 Intro Narration
    headline = article_text.split("\n")[0].strip().replace('"', '')

    if content_language == "Hindi":
        slide1_limit = slide_char_limits.get(1, SLIDE_CHAR_LIMITS[1])
        slide1_prompt = (
            f"Generate news headline narration in Hindi for the story: {headline}. "
            f"Maximum {slide1_limit} characters. Avoid greetings. Respond in Hindi (Devanagari script) only."
        )
    else:
        slide1_limit = slide_char_limits.get(1, SLIDE_CHAR_LIMITS[1])
        slide1_prompt = (
            f"Generate headline intro narration in English for: {headline}. "
            f"Maximum {slide1_limit} characters. Avoid greetings. Respond in English only, translating the source if necessary."
        )

    slide1_response = client.chat.completions.create(
        model="gpt-5-chat",
        messages=[
            {"role": "system", "content": "You are a news presenter generating opening lines."},
            {"role": "user", "content": slide1_prompt}
        ]
    )
    slide1_script = textwrap.shorten(
        slide1_response.choices[0].message.content.strip(),
        width=slide1_limit,
        placeholder="…"
    )

    slides = [{
        "title": headline[:80],
        "prompt": "Intro slide with headline framing.",
        "image_prompt": f"Vector-style illustration of Polaris presenting news: {headline}",
        "script": slide1_script
    }]

    # 🔹 Generate narration for each slide
    for slide in slides_raw:
        script_language = f"{content_language} (use Devanagari script)" if content_language == "Hindi" else content_language
        story_slide_index = len(slides) + 1  # includes intro already
        target_limit = slide_char_limits.get(story_slide_index, default_limit)
        language_requirement = (
            "Deliver the narration strictly in Hindi (Devanagari script)."
            if content_language == "Hindi"
            else "Deliver the narration strictly in English. Do not include Hindi words or transliteration."
        )

        caption = (slide.get("title") or "").strip()
        summary_brief = (slide.get("summary") or slide.get("caption") or slide.get("prompt") or "").strip()
        image_prompt = (slide.get("image_prompt") or "").strip()

        if not summary_brief:
            summary_brief = caption or "Provide factual narration for this segment."

        narration_prompt = f"""
Write a narration in **{script_language}** (max {target_limit} characters),
in the voice of Polaris (factual, vivid, and neutral). {language_requirement}

Key points to cover:
{summary_brief}

Visual inspiration:
{image_prompt or 'Use a neutral newsroom-inspired background.'}

Character sketch:
{character_sketch}
"""

        try:
            narration_response = client.chat.completions.create(
                model="gpt-5-chat",
                messages=[
                    {"role": "system", "content": "You write concise narrations for web story slides."},
                    {"role": "user", "content": narration_prompt.strip()}
                ]
            )
            narration = textwrap.shorten(
                narration_response.choices[0].message.content.strip(),
                width=target_limit,
                placeholder="…"
            )
        except:
            narration = "Unable to generate narration for this slide."

        slides.append({
            "title": caption or slide.get("title", ""),
            "prompt": summary_brief,
            "image_prompt": image_prompt or f"Modern vector-style visual for: {caption or slide.get('title', '')}",
            "script": narration
        })

    return {
        "category": category,
        "subcategory": subcategory,
        "emotion": emotion,
        "slides": slides
    }



def modify_tab4_json(original_json):
    updated_json = OrderedDict()
    slide_number = 2  # Start from slide2 since slide1 & slide2 are removed

    for i in range(3, 100):  # Covers slide3 to slide99
        old_key = f"slide{i}"
        if old_key not in original_json:
            break
        content = original_json[old_key]
        new_key = f"slide{slide_number}"

        for k, v in content.items():
            if k.endswith("paragraph1"):
                para_key = f"s{slide_number}paragraph1"
                audio_key = f"audio_url{slide_number}"
                updated_json[new_key] = {
                    para_key: v,
                    audio_key: content.get("audio_url", ""),
                    "voice": content.get("voice", "")
                }
                break
        slide_number += 1

    return updated_json

# Function to generate an AMP slide using paragraph, audio URL, and background image
def generate_slide(paragraph: str, audio_url: str, background_image_url: str = None):
    """Generate AMP slide with optional background image"""
    # Default background image if none provided
    if not background_image_url:
        background_image_url = "https://media.suvichaar.org/upload/polaris/polarisslide.png"
    
    return f"""
        <amp-story-page id="cover-slide" auto-advance-after="cover-audio">
        <amp-story-grid-layer template="fill">
          <amp-img src="{background_image_url}"
            width="720" height="1280" layout="responsive">
          </amp-img>
        </amp-story-grid-layer>
        <amp-story-grid-layer template="fill">
          <amp-video autoplay loop layout="fixed" width="1" height="1" poster="" id="cover-audio">
            <source type="audio/mpeg" src="{audio_url}">
          </amp-video>
        </amp-story-grid-layer>
        <amp-story-grid-layer template="vertical">
          <div class="centered-container">
            
            <div class="text1">
              {paragraph}
            </div>
           <div class="footer"><p>©SuvichaarAI</p></div>
          </div>
        </amp-story-grid-layer>
      </amp-story-page>
        """

def replace_placeholders_in_html(html_text, json_data, background_image_url=None):
    storytitle = json_data.get("slide1", {}).get("storytitle", "")
    storytitle_url = json_data.get("slide1", {}).get("audio_url", "")

    html_text = html_text.replace("{{storytitle}}", storytitle)
    html_text = html_text.replace("{{storytitle_audiourl}}", storytitle_url)

    # Replace hardcoded images in template with user's uploaded image (if provided)
    if background_image_url:
        html_text = html_text.replace(
            "https://media.suvichaar.org/upload/polaris/polariscover.png",
            background_image_url
        )
        html_text = html_text.replace(
            "https://media.suvichaar.org/upload/polaris/polarisslide.png",
            background_image_url
        )

    # Insert middle slides (skip slide1 which is storytitle)
    all_slides = ""
    for key in sorted(
        json_data.keys(),
        key=lambda x: int(x.replace("slide", "")) if x.startswith("slide") else 999999
    ):
        if not key.startswith("slide"):
            continue

        slide_num = int(key.replace("slide", ""))
        if slide_num == 1:
            continue

        data = json_data[key]
        if not isinstance(data, dict):
            continue

        paragraph = ""
        audio_url = data.get("audio_url", "")
        for k, v in data.items():
            if k.startswith("s") and "paragraph1" in k and isinstance(v, str):
                paragraph = v.replace("'", "'").replace('"', '&quot;')
                break

        if paragraph and audio_url:
            all_slides += generate_slide(paragraph, audio_url, background_image_url)

    html_text = html_text.replace("<!--INSERT_SLIDES_HERE-->", all_slides)

    return html_text

def restructure_slide_output(final_output, limits):
    slides = final_output.get("slides", [])
    structured = {}
    default_limit = limits.get("default", 200)

    for idx, slide in enumerate(slides):
        key = f"s{idx + 1}paragraph1"
        script = slide.get("script", "").strip()

        # Safety net: If empty script, fall back to title or prompt
        if not script:
            fallback = slide.get("title") or slide.get("prompt") or "Content unavailable"
            script = fallback.strip()

        target_limit = limits.get(idx + 1, default_limit)
        if len(script) > target_limit:
            script = textwrap.shorten(script, width=target_limit, placeholder="…")

        structured[key] = script

    return structured

def generate_remotion_input(tts_output: dict, fixed_image_url: str, author_name: str = "Suvichaar"):
    remotion_data = OrderedDict()
    slide_index = 1

    # Slide 1: storytitle
    if "storytitle" in tts_output:
        remotion_data[f"slide{slide_index}"] = {
            f"s{slide_index}paragraph1": tts_output["storytitle"],
            f"s{slide_index}audio1": tts_output.get(f"slide{slide_index}", {}).get("audio_url", ""),
            f"s{slide_index}image1": fixed_image_url,
            f"s{slide_index}paragraph2": f"- {author_name}"
        }
        slide_index += 1

    # Slides for s1paragraph1 to s9paragraph1
    for i in range(1, 10):
        key = f"s{i}paragraph1"
        if key in tts_output:
            slide_key = f"slide{slide_index}"
            remotion_data[slide_key] = {
                f"s{slide_index}paragraph1": tts_output[key],
                f"s{slide_index}audio1": tts_output.get(slide_key, {}).get("audio_url", ""),
                f"s{slide_index}image1": fixed_image_url,
                f"s{slide_index}paragraph2": f"- {author_name}"
            }
            slide_index += 1

    # ✅ Final CTA slide (always last)
    remotion_data[f"slide{slide_index}"] = {
        f"s{slide_index}paragraph1": "For Such Content\nStay Connected with Suvichar Live",
        f"s{slide_index}audio1": DEFAULT_CTA_AUDIO,
        f"s{slide_index}video1": "",
        f"s{slide_index}paragraph2": "Read|Share|Inspire"
    }

    # Save to file
    timestamp = int(time.time())
    filename = f"remotion_input_{timestamp}.json"
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(remotion_data, f, indent=2, ensure_ascii=False)

    return filename

# ------------------ Azure Speech Neural TTS Helper ------------------
def azure_tts_generate(text: str, voice: str, retries: int = 2, backoff: float = 1.0) -> bytes:
    """
    Generate speech bytes using Azure Speech SDK neural voices.
    Uses region-based configuration instead of endpoint.
    """
    speech_config = speechsdk.SpeechConfig(
        subscription=AZURE_SPEECH_KEY,
        region=AZURE_SPEECH_REGION
    )
    speech_config.speech_synthesis_voice_name = voice
    speech_config.set_speech_synthesis_output_format(
        speechsdk.SpeechSynthesisOutputFormat.Audio16Khz32KBitRateMonoMp3
    )

    # ✅ Output to memory — safe for Streamlit Cloud
    audio_config = None

    for attempt in range(retries + 1):
        synthesizer = speechsdk.SpeechSynthesizer(
            speech_config=speech_config,
            audio_config=audio_config
        )
        result = synthesizer.speak_text_async(text).get()

        if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
            return result.audio_data  # Return bytes directly

        if result.reason == speechsdk.ResultReason.Canceled and attempt < retries:
            time.sleep(backoff * (2 ** attempt))
            continue

        if result.reason == speechsdk.ResultReason.Canceled:
            details = result.cancellation_details
            raise RuntimeError(
                f"Azure TTS canceled: {details.reason}; error={getattr(details, 'error_details', None)}"
            )
        else:
            raise RuntimeError(f"Azure TTS failed with reason: {result.reason}")

    raise RuntimeError("Azure TTS failed after retries")


def synthesize_and_upload(paragraphs, voice):
    s3 = boto3.client(
        "s3",
        aws_access_key_id=AWS_ACCESS_KEY,
        aws_secret_access_key=AWS_SECRET_KEY,
        region_name=AWS_REGION,
    )

    result = OrderedDict()
    os.makedirs("temp", exist_ok=True)

    slide_index = 1

    # Slide 1: storytitle
    if "storytitle" in paragraphs:
        storytitle = paragraphs["storytitle"]
        audio_bytes = azure_tts_generate(storytitle, voice)
        filename = f"tts_{uuid.uuid4().hex}.mp3"
        local_path = os.path.join("temp", filename)
        with open(local_path, "wb") as f:
            f.write(audio_bytes)
        s3_key = f"{S3_PREFIX}{filename}"
        s3.upload_file(local_path, AWS_BUCKET, s3_key)
        cdn_url = f"{CDN_BASE}{s3_key}"
        result[f"slide{slide_index}"] = {
            "storytitle": storytitle,
            "audio_url": cdn_url,
            "voice": voice
        }
        os.remove(local_path)
        slide_index += 1

    # Slide 2..(N-1) : s1paragraph1.. 
    for i in range(1, 50):
        key = f"s{i}paragraph1"
        if key not in paragraphs:
            break
        text_val = paragraphs[key]
        st.write(f"Processing {key}")
        audio_bytes = azure_tts_generate(text_val, voice)
        filename = f"tts_{uuid.uuid4().hex}.mp3"
        local_path = os.path.join("temp", filename)
        with open(local_path, "wb") as f:
            f.write(audio_bytes)
        s3_key = f"{S3_PREFIX}{filename}"
        s3.upload_file(local_path, AWS_BUCKET, s3_key)
        cdn_url = f"{CDN_BASE}{s3_key}"
        result[f"slide{slide_index}"] = {key: text_val, "audio_url": cdn_url, "voice": voice}
        os.remove(local_path)
        slide_index += 1

    return result

def transliterate_to_devanagari(json_data):
    updated = {}

    for k, v in json_data.items():
        # Only transliterate slide paragraphs
        if k.startswith("s") and "paragraph1" in k and v.strip():
            prompt = f"""Transliterate this Hindi sentence (written in Latin script) into Hindi Devanagari script. Return only the transliterated text:\n\n{v}"""
            
            try:
                response = client.chat.completions.create(
                    model="gpt-5-chat",
                    messages=[
                        {"role": "system", "content": "You are a Hindi transliteration expert."},
                        {"role": "user", "content": prompt.strip()}
                    ]
                )
                devanagari = response.choices[0].message.content.strip()
                updated[k] = devanagari
            except Exception as e:
                # Fallback: use original if error occurs
                updated[k] = v
        else:
            updated[k] = v

    return updated

def generate_storytitle(title, summary, content_language="English", max_chars=180):
    if content_language == "Hindi":
        prompt = f"""
आप एक समाचार शीर्षक विशेषज्ञ हैं। नीचे दी गई अंग्रेज़ी समाचार शीर्षक और सारांश को पढ़कर, उसी का अर्थ बनाए रखते हुए एक नया आकर्षक **हिंदी शीर्षक** बनाइए।

अंग्रेज़ी शीर्षक: {title}
सारांश: {summary}

अनुरोध:
- केवल एक वाक्य
- भाषा सरल और स्पष्ट हो
- भावनात्मक रूप से आकर्षक हो
- उद्धरण ("") शामिल न करें
- अधिकतम {max_chars} वर्ण

अब कृपया हिंदी शीर्षक दीजिए:
"""
    else:
        prompt = f"""
You are an experienced news editor. Craft a single-sentence English headline (maximum {max_chars} characters) that captures the essence of the following story. 
Highlight the key actors and the central development. Avoid quotes, hashtags, or greetings.

Original title: {title}
Summary: {summary}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-5-chat",
            messages=[
                {"role": "system", "content": "You generate clear and concise news headlines."},
                {"role": "user", "content": prompt.strip()}
            ]
        )
        raw_title = response.choices[0].message.content.strip().strip('"')
        return textwrap.shorten(raw_title, width=max_chars, placeholder="…")

    except Exception as e:
        print(f"❌ Storytitle generation failed: {e}")
        fallback = title.strip()
        if summary:
            fallback = f"{title.strip()}: {summary.strip()}"
        return textwrap.shorten(fallback, width=max_chars, placeholder="…")


# === Streamlit UI ===
st.title("🧠 Web Story Content Generator")

# Input options
st.subheader("📥 Input Options")

input_method = st.radio(
    "Choose input method:",
    ["Article URL", "Text Field", "Upload PDF", "Upload DOCX", "Upload Image (OCR)"],
    horizontal=True
)

url = None
manual_text = None
uploaded_file = None
article_image_file = None

if input_method == "Article URL":
    url = st.text_input("Enter a news article URL")
elif input_method == "Text Field":
    manual_text = st.text_area("Enter article text", height=200)
    if manual_text:
        # Extract title and summary from manual text
        lines = manual_text.strip().split('\n')
        url = None  # Will be handled separately
elif input_method == "Upload PDF":
    uploaded_file = st.file_uploader("Upload PDF file", type=["pdf"])
    if uploaded_file:
        url = None
elif input_method == "Upload DOCX":
    uploaded_file = st.file_uploader("Upload DOCX file", type=["docx"])
    if uploaded_file:
        url = None
elif input_method == "Upload Image (OCR)":
    article_image_file = st.file_uploader("Upload an article image (JPG/PNG) for OCR", type=["jpg", "jpeg", "png", "webp"])
    if article_image_file:
        url = None

# Image upload
st.subheader("🖼️ Cover Image (Optional)")
uploaded_image = st.file_uploader(
    "Upload cover image (will replace default background)",
    type=["jpg", "jpeg", "png", "webp"]
)

# Optional: Extract text from image using OCR
extract_text_from_uploaded_image = False
if uploaded_image:
    extract_text_from_uploaded_image = st.checkbox(
        "Extract text from this image using OCR (Azure Document Intelligence)",
        value=False,
        help="If enabled, will extract text from the uploaded image and use it as article content"
    )

# Output language selection
st.subheader("🌐 Output Settings")
content_language = st.selectbox(
    "Output Language",
    ["English", "Hindi"],
    index=0
)

# Number of slides
persona = "Expert news anchor"  # Fixed persona
number = st.number_input(
    "Enter total number of slides (including Storytitle and CTA slide)",
    min_value=4,
    max_value=10,
    value=8,
    step=1
)
# User enters total slides, we need middle slides (minus 2 for storytitle and CTA)
middle_count = max(2, number - 2)

# Character limit per slide
max_chars = st.slider(
    "Max characters per slide",
    min_value=150,
    max_value=260,
    value=200,
    step=10
)

if st.button("🚀 Generate Complete Web Story"):
    # Determine which input method was used and extract content
    title = None
    summary = None
    full_text = None
    
    if input_method == "Article URL" and url:
        with st.spinner("🔄 Step 1/4: Analyzing article..."):
            try:
                title, summary, full_text = extract_article(url)
            except Exception as e:
                st.error(f"❌ Error extracting article: {str(e)}")
                st.stop()
    elif input_method == "Text Field" and manual_text:
        with st.spinner("🔄 Step 1/4: Processing text..."):
            try:
                lines = manual_text.strip().split('\n')
                title = lines[0] if lines else "Untitled Article"
                summary = manual_text[:300].strip()
                full_text = manual_text.strip()
            except Exception as e:
                st.error(f"❌ Error processing text: {str(e)}")
                st.stop()
    elif input_method == "Upload PDF" and uploaded_file:
        with st.spinner("🔄 Step 1/4: Extracting from PDF..."):
            try:
                title, summary, full_text = extract_from_pdf(uploaded_file)
                if not title:
                    st.error("❌ Could not extract content from PDF")
                    st.stop()
            except Exception as e:
                st.error(f"❌ Error extracting from PDF: {str(e)}")
                st.stop()
    elif input_method == "Upload DOCX" and uploaded_file:
        with st.spinner("🔄 Step 1/4: Extracting from DOCX..."):
            try:
                title, summary, full_text = extract_from_docx(uploaded_file)
                if not title:
                    st.error("❌ Could not extract content from DOCX")
                    st.stop()
            except Exception as e:
                st.error(f"❌ Error extracting from DOCX: {str(e)}")
                st.stop()
    elif input_method == "Upload Image (OCR)" and article_image_file:
        with st.spinner("🔄 Step 1/4: Extracting text from article image (OCR)..."):
            try:
                ocr_text = extract_text_from_image(article_image_file)
                if ocr_text:
                    lines = ocr_text.strip().split('\n')
                    title = lines[0] if lines else "Untitled from Image"
                    summary = ocr_text[:300].strip()
                    full_text = ocr_text.strip()
                else:
                    st.error("❌ Could not extract text from the uploaded article image")
                    st.stop()
            except Exception as e:
                st.error(f"❌ Error extracting text from article image: {str(e)}")
                st.stop()
    elif extract_text_from_uploaded_image and uploaded_image:
        with st.spinner("🔄 Step 1/4: Extracting text from image (OCR)..."):
            try:
                ocr_text = extract_text_from_image(uploaded_image)
                if ocr_text:
                    lines = ocr_text.strip().split('\n')
                    title = lines[0] if lines else "Untitled from Image"
                    summary = ocr_text[:300].strip()
                    full_text = ocr_text.strip()
                else:
                    st.error("❌ Could not extract text from image")
                    st.stop()
            except Exception as e:
                st.error(f"❌ Error extracting text from image: {str(e)}")
                st.stop()
    else:
        if not (input_method == "Article URL" and url) and not (input_method == "Text Field" and manual_text) and not (input_method == "Upload PDF" and uploaded_file) and not (input_method == "Upload DOCX" and uploaded_file):
            st.warning("⚠️ Please provide input content (URL, text, PDF, DOCX, or enable OCR on uploaded image)")
            st.stop()
    
    if title and summary and full_text:
        with st.spinner("🔄 Step 1/4: Analyzing content..."):
            try:
                sentiment = get_sentiment(summary or full_text)
                result = detect_category_and_subcategory(full_text, content_language)
                category = result["category"]
                subcategory = result["subcategory"]
                emotion = result["emotion"]

                slide_char_limits = SLIDE_CHAR_LIMITS.copy()
                slide_char_limits["default"] = max_chars

                storytitle_language = "Hindi" if content_language == "Hindi" else "English"
                storytitle = generate_storytitle(
                    title,
                    summary,
                    storytitle_language,
                    max_chars=slide_char_limits.get(1, SLIDE_CHAR_LIMITS[1])
                )
                
                st.session_state.story_title_from_tab1 = storytitle

                # Step 7: Generate slide content
                output = title_script_generator(
                    category,
                    subcategory,
                    emotion,
                    full_text,
                    content_language,
                    middle_count=middle_count,
                    slide_char_limits=slide_char_limits
                )

                final_output = {
                    "title": title,
                    "summary": summary,
                    "sentiment": sentiment,
                    "emotion": emotion,
                    "category": category,
                    "subcategory": subcategory,
                    "persona": persona,
                    "slides": output.get("slides", []),
                    "storytitle": storytitle
                }

                # Step 8: Flatten into story JSON
                structured_output = OrderedDict()
                structured_output["storytitle"] = textwrap.shorten(
                    storytitle,
                    width=slide_char_limits.get(1, SLIDE_CHAR_LIMITS[1]),
                    placeholder="…"
                )

                flattened_slides = restructure_slide_output(final_output, slide_char_limits)
                for key, value in flattened_slides.items():
                    structured_output[key] = value
                
            except Exception as e:
                st.error(f"❌ Error in content generation: {str(e)}")
                st.stop()
            
            # Step 2: TTS
            with st.spinner("🔄 Step 2/4: Generating audio..."):
                try:
                    # Select voice based on language
                    if content_language == "Hindi":
                        voice_label = "hi-IN-AaravNeural"
                    else:
                        voice_label = "en-IN-AaravNeural"
                    tts_output = synthesize_and_upload(structured_output, voice_label)
                except Exception as e:
                    st.error(f"❌ Error in TTS: {str(e)}")
                    st.stop()
            
            # Initialize image variables (needed across steps)
            original_image_url = DEFAULT_COVER_URL
            resized_image_url = None
            image_s3_key = None
            
            # Step 3: HTML Generation
            with st.spinner("🔄 Step 3/4: Building HTML..."):
                try:
                    # Handle image upload first (needed for slide backgrounds)
                    
                    if uploaded_image:
                        with st.spinner("Uploading image to S3..."):
                            original_image_url, image_s3_key = upload_image_to_s3(uploaded_image, bucket_name=AWS_BUCKET)
                            if original_image_url and image_s3_key:
                                # Generate resized image URL for slide backgrounds (412x618 for web story aspect ratio)
                                resized_image_url = get_resized_image_url(image_s3_key, width=720, height=1280, bucket_name=AWS_BUCKET)
                                if not resized_image_url:
                                    st.warning("⚠️ Could not generate resized image, using default")
                                    resized_image_url = None
                            else:
                                original_image_url = DEFAULT_COVER_URL
                                st.warning("⚠️ Using default image due to upload error")
                    
                    with open("templates/test.html", "r", encoding="utf-8") as f:
                        html_template = f.read()
                    
                    # Pass resized image URL for slide backgrounds
                    updated_html = replace_placeholders_in_html(html_template, tts_output, background_image_url=resized_image_url)
                    
                except Exception as e:
                    st.error(f"❌ Error in HTML generation: {str(e)}")
                    st.stop()
            
            # Step 4: Final HTML with meta
            with st.spinner("🔄 Step 4/4: Adding meta tags..."):
                try:
                    # Use original image URL (non-resized) for image0 placeholder
                    image_url = original_image_url
                    
                    # Generate metadata
                    messages = [{
                        "role": "user",
                        "content": f"""
                        Generate the following for a web story titled '{storytitle}':
                        1. A short SEO-friendly meta description
                        2. Meta keywords (comma separated)
                        3. Relevant filter tags (comma separated, suitable for categorization and content filtering)
                        4. Content Type (choose ONLY one: "News" or "Article")
                        5. Primary Language (choose ONLY one: "en-US" or "hi-IN")
                        6. Category (choose ONLY one from: Art, Travel, Entertainment, Literature, Books, Sports, History, Culture, Wildlife, Spiritual, Food)
                        
                        Format your response as:
                        Description: <meta description>
                        Keywords: <keywords>
                        Filter Tags: <tags>
                        Content Type: <News or Article>
                        Language: <en-US or hi-IN>
                        Category: <one of the categories listed>"""
                    }]
                    
                    response = client.chat.completions.create(
                        model="gpt-5-chat",
                        messages=messages,
                        max_tokens=300,
                        temperature=0.5,
                    )
                    output = response.choices[0].message.content
            
                    # Extract metadata
                    desc = re.search(r"[Dd]escription\s*[:\-]\s*(.+)", output)
                    keys = re.search(r"[Kk]eywords\s*[:\-]\s*(.+)", output)
                    content_type_match = re.search(r"[Cc]ontent\s*[Tt]ype\s*[:\-]\s*(News|Article)", output)
                    lang_match = re.search(r"[Ll]anguage\s*[:\-]\s*(en-US|hi-IN)", output)
            
                    meta_description = desc.group(1).strip() if desc else ""
                    meta_keywords = keys.group(1).strip() if keys else ""
                    content_type = content_type_match.group(1) if content_type_match else "News"
                    language = lang_match.group(1) if lang_match else ("hi-IN" if content_language == "Hindi" else "en-US")
                    
                    # Generate URLs
                    nano, slug_nano, canurl, canurl1 = generate_slug_and_urls(storytitle)
                    page_title = f"{storytitle}| Suvichaar"
                    
                    # User mapping
                    user_mapping = {
                        "Mayank": "https://www.instagram.com/iamkrmayank?igsh=eW82NW1qbjh4OXY2&utm_source=qr",
                        "Onip": "https://www.instagram.com/onip.mathur/profilecard/?igsh=MW5zMm5qMXhybGNmdA==",
                        "Naman": "https://njnaman.in/"
                    }
                    selected_user = random.choice(list(user_mapping.keys()))
                    
                    # Replace all placeholders - do this BEFORE image URL replacements
                    updated_html = updated_html.replace("{{user}}", selected_user)
                    updated_html = updated_html.replace("{{userprofileurl}}", user_mapping[selected_user])
                    updated_html = updated_html.replace("{{publishedtime}}", datetime.now(timezone.utc).isoformat(timespec='seconds'))
                    updated_html = updated_html.replace("{{modifiedtime}}", datetime.now(timezone.utc).isoformat(timespec='seconds'))
                    updated_html = updated_html.replace("{{pagetitle}}", page_title)
                    updated_html = updated_html.replace("{{canurl}}", canurl)
                    updated_html = updated_html.replace("{{canurl1}}", canurl1)
                    updated_html = updated_html.replace("{{metadescription}}", meta_description)
                    updated_html = updated_html.replace("{{metakeywords}}", meta_keywords)
                    updated_html = updated_html.replace("{{contenttype}}", content_type)
                    updated_html = updated_html.replace("{{lang}}", language)
                    updated_html = updated_html.replace("{{storytitle}}", storytitle)
                    updated_html = updated_html.replace("{{category}}", category)
                    
                    # Replace storytitle audio URL
                    storytitle_audio = tts_output.get("slide1", {}).get("audio_url", "")
                    updated_html = updated_html.replace("{{storytitle_audiourl}}", storytitle_audio)
                    
                    # Replace site logo placeholders with default values
                    default_logo_base = "https://media.suvichaar.org/filters:resize"
                    updated_html = updated_html.replace("{{sitelogo32x32}}", f"{default_logo_base}/32x32/media/brandasset/suvichaariconblack.png")
                    updated_html = updated_html.replace("{{sitelogo192x192}}", f"{default_logo_base}/192x192/media/brandasset/suvichaariconblack.png")
                    updated_html = updated_html.replace("{{sitelogo180x180}}", f"{default_logo_base}/180x180/media/brandasset/suvichaariconblack.png")
                    updated_html = updated_html.replace("{{sitelogo144x144}}", f"{default_logo_base}/144x144/media/brandasset/suvichaariconblack.png")
                    updated_html = updated_html.replace("{{sitelogo96x96}}", f"{default_logo_base}/96x96/media/brandasset/suvichaariconblack.png")
                    
                    # Replace organization placeholder
                    updated_html = updated_html.replace("{{organization}}", "Suvichaar")
                    
                    # Replace publisher placeholder
                    updated_html = updated_html.replace("{{publisher}}", "Suvichaar")
                    updated_html = updated_html.replace("{{publisherlogosrc}}", "https://media.suvichaar.org/media/designasset/brandasset/icons/quaternary/whitequaternaryicon.png")
                    
                    # Replace prev/next story placeholders (empty for now)
                    updated_html = updated_html.replace("{{prevstorytitle}}", "")
                    updated_html = updated_html.replace("{{prevstorylink}}", "")
                    updated_html = updated_html.replace("{{nextstorytitle}}", "")
                    updated_html = updated_html.replace("{{nextstorylink}}", "")
                    
                    # Use original image URL (non-resized) for image0 placeholder
                    updated_html = updated_html.replace("{{image0}}", image_url)
                    
                    # Parse the image URL to get the path for resizing (for portrait and thumbnail)
                    parsed_cdn_url = urlparse(image_url)
                    cdn_key_path = parsed_cdn_url.path.lstrip("/")
                    
                    # If image is from CDN, extract the S3 key
                    if CDN_BASE in image_url:
                        cdn_key_path = image_url.replace(CDN_BASE, "").lstrip("/")
                    elif image_s3_key:
                        # Use the S3 key we have from upload
                        cdn_key_path = image_s3_key
                    
                    resize_presets = {
                        "potraitcoverurl": (720, 1280),
                        "msthumbnailcoverurl": (300, 300),
                    }
                    
                    for label, (width, height) in resize_presets.items():
                        template = {
                            "bucket": AWS_BUCKET,
                            "key": cdn_key_path,
                            "edits": {
                                "resize": {
                                    "width": width,
                                    "height": height,
                                    "fit": "cover"
                                }
                            }
                        }
                        encoded = base64.urlsafe_b64encode(json.dumps(template).encode()).decode()
                        final_url = f"{CDN_PREFIX_MEDIA}{encoded}"
                        # Replace placeholder - ensure no curly braces in replacement
                        placeholder = f"{{{{{label}}}}}"
                        updated_html = updated_html.replace(placeholder, final_url)
                    
                    # Comprehensive cleanup: Remove curly braces from URLs anywhere in the HTML
                    # This handles cases where URLs might have been wrapped in braces
                    updated_html = re.sub(r'\{(\s*https?://[^\s"\'<>}]+)\}', r'\1', updated_html)
                    updated_html = re.sub(r'(\s*https?://[^\s"\'<>}]+)\}', r'\1', updated_html)
                    updated_html = re.sub(r'\{(\s*https?://[^\s"\'<>}]+)', r'\1', updated_html)
                    
                    # Also clean up in attribute values (href, src, content, etc.)
                    updated_html = re.sub(r'(href|src|content|url|poster-portrait-src|publisher-logo-src)="\{([^"]+)\}"', r'\1="\2"', updated_html)
                    updated_html = re.sub(r'(href|src|content|url|poster-portrait-src|publisher-logo-src)=\'\{([^\']+)\}\'', r'\1=\'\2\'', updated_html)
                    
                    # Clean up JSON-LD script content (URLs in JSON strings)
                    updated_html = re.sub(r'"\{([^"]+)\}"', r'"\1"', updated_html)
                    updated_html = re.sub(r':\s*"\{([^"]+)\}"', r': "\1"', updated_html)
                    
                    # Final comprehensive cleanup: Remove any remaining curly braces around URLs
                    # This catches any edge cases we might have missed
                    # Pattern: {https://...} or {https://... or https://...}
                    updated_html = re.sub(r'\{(\s*https?://[^}\s"\'<>]+)\}', r'\1', updated_html)
                    updated_html = re.sub(r'\{(\s*https?://[^}\s"\'<>]+)', r'\1', updated_html)
                    updated_html = re.sub(r'(\s*https?://[^}\s"\'<>]+)\}', r'\1', updated_html)
                    
                    # Remove curly braces from JSON values (more specific pattern)
                    updated_html = re.sub(r':\s*"\{([^"]+)\}"', r': "\1"', updated_html)
                    updated_html = re.sub(r':\s*"([^"]*)\{([^"]+)\}([^"]*)"', r': "\1\2\3"', updated_html)
                    
                    # Upload HTML file to S3 bucket "suvichaarstories"
                    html_filename = f"{slug_nano}.html"
                    html_s3_key = html_filename
                    html_bucket = "suvichaarstories"
                    
                    try:
                        with st.spinner("📤 Uploading HTML file to S3..."):
                            s3_client.put_object(
                                Bucket=html_bucket,
                                Key=html_s3_key,
                                Body=updated_html.encode('utf-8'),
                                ContentType='text/html'
                            )
                        st.success(f"✅ HTML file uploaded successfully to S3!")
                    except Exception as e:
                        st.warning(f"⚠️ Failed to upload HTML to S3: {str(e)}")
                    
                    # Display canurl
                    st.info(f"🔗 **Story URL:** {canurl}")
                    st.info(f"🔗 **Story URL (Alternative):** {canurl1}")
                    
                    st.success("✅ Complete! Your web story is ready!")
                    
                    # Download HTML file
                    st.download_button(
                        label="⬇️ Download Final HTML",
                        data=updated_html,
                        file_name=html_filename,
                        mime="text/html"
                    )
                    
                except Exception as e:
                    st.error(f"❌ Error in finalization: {str(e)}")
    else:
        st.warning("Please enter a valid URL.")
